"""
Document Ingestion Service - Process documents and add to FAISS KB (no Supabase)
"""

import os
import hashlib
from typing import List, Dict, Optional, Tuple
from pathlib import Path
from datetime import datetime
import uuid
import numpy as np

try:
    import PyPDF2
    import pdfplumber
    from docx import Document as DocxDocument
    import openpyxl
except ImportError:
    pass

try:
    import faiss
    from sentence_transformers import SentenceTransformer
except ImportError:
    pass


class DocumentIngestionService:
    """
    Service for ingesting documents into FAISS knowledge base (local storage only).
    
    Supports:
    - PDF files (with text extraction)
    - DOCX files
    - XLSX files
    - Plain text files
    
    Stores vectors in FAISS index + metadata in SQLite (no Supabase)
    """
    
    CHUNK_SIZE = 512
    CHUNK_OVERLAP = 50
    
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50, 
                 retriever=None):
        """
        Initialize ingestion service.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
            retriever: RetrieverService instance for adding to FAISS
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.retriever = retriever
        self.processed_docs = 0
        self.failed_docs = 0
        self.total_chunks_created = 0
        
        # Initialize embedding model for chunk encoding
        self.embedding_model = None
        self._load_embedding_model()
    
    def _load_embedding_model(self):
        """Load embedding model for chunk encoding"""
        try:
            os.environ.setdefault("HF_HUB_OFFLINE", "1")
            self.embedding_model = SentenceTransformer("all-MiniLM-L6-v2")
        except:
            print("Warning: Embedding model not available, chunks will use fallback encoding")
    
    def ingest_file(self, file_path: str, source: str = "manual") -> Dict:
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {
                "status": "error",
                "message": f"File not found: {file_path}",
                "chunks": []
            }
        
        file_ext = file_path.suffix.lower()
        
        try:
            if file_ext == ".pdf":
                chunks = self._extract_pdf(str(file_path))
            elif file_ext == ".docx":
                chunks = self._extract_docx(str(file_path))
            elif file_ext == ".xlsx":
                chunks = self._extract_xlsx(str(file_path))
            elif file_ext == ".txt":
                chunks = self._extract_text(str(file_path))
            else:
                return {
                    "status": "error",
                    "message": f"Unsupported file type: {file_ext}",
                    "chunks": []
                }
            
            # Create chunk objects with metadata
            processed_chunks = self._create_chunks(
                chunks=chunks,
                filename=file_path.name,
                source=source
            )
            
            # Add chunks to FAISS index if retriever is available
            if self.retriever and processed_chunks:
                add_result = self.retriever.add_chunks(processed_chunks)
                print(f"FAISS index update: {add_result['message']}")
            
            self.processed_docs += 1
            
            return {
                "status": "success",
                "filename": file_path.name,
                "chunks_created": len(processed_chunks),
                "chunks": processed_chunks,
                "source": source
            }
        
        except Exception as e:
            self.failed_docs += 1
            return {
                "status": "error",
                "message": f"Failed to ingest file: {str(e)}",
                "chunks": []
            }
    
    def _extract_pdf(self, file_path: str) -> List[str]:
        """Extract text from PDF file"""
        text_parts = []
        
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        
        return text_parts
    
    def _extract_docx(self, file_path: str) -> List[str]:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return text_parts
    
    def _extract_xlsx(self, file_path: str) -> List[str]:
        """Extract text from XLSX file"""
        wb = openpyxl.load_workbook(file_path)
        text_parts = []
        
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            for row in ws.iter_rows(values_only=True):
                row_text = " | ".join(str(cell) for cell in row if cell)
                if row_text.strip():
                    text_parts.append(row_text)
        
        return text_parts
    
    def _extract_text(self, file_path: str) -> List[str]:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        return [content]
    
    def _create_chunks(
        self,
        chunks: List[str],
        filename: str,
        source: str
    ) -> List[Dict]:
        """
        Create overlapping chunks from extracted text.
        
        Args:
            chunks: List of text chunks from document
            filename: Source filename
            source: Source identifier
            
        Returns:
            List of chunk dictionaries with metadata
        """
        all_chunks = []
        full_text = "\n".join(chunks)
        
        # Split into overlapping chunks
        for i in range(0, len(full_text), self.chunk_size - self.chunk_overlap):
            chunk_text = full_text[i:i + self.chunk_size]
            
            if len(chunk_text.strip()) > 50:  # Minimum chunk size
                chunk_id = str(uuid.uuid4())
                all_chunks.append({
                    "chunk_id": chunk_id,
                    "doc_id": f"doc-{hashlib.md5(filename.encode()).hexdigest()[:8]}",
                    "text": chunk_text,
                    "source_file": filename,
                    "source": source,
                    "approval_status": "pending",
                    "document_version": 1,
                    "created_at": datetime.utcnow().isoformat() + "Z",
                    "chunk_index": len(all_chunks)
                })
        
        return all_chunks
    
    def get_stats(self) -> Dict:
        """Get ingestion statistics"""
        return {
            "processed_documents": self.processed_docs,
            "failed_documents": self.failed_docs,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap
        }
